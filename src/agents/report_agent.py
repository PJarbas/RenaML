"""ReportAgent: Generates final reports in multiple formats with LLM executive summary."""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

import markdown
from agno.agent import Agent
from agno.models.openai import OpenAIChat
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from prompts.report_agent_prompts import AGENT_INSTRUCTIONS, EXECUTIVE_SUMMARY_PROMPT_TEMPLATE


class ReportAgent:
    """Agent responsible for generating final reports with LLM-powered executive summary."""

    def __init__(self, team_memory: dict[str, Any], run_dir: Path) -> None:
        self.team_memory = team_memory
        self.run_dir = run_dir
        self.logger = logging.getLogger(__name__)

        # Create Agent instance for LLM executive summary
        self.llm_agent = Agent(
            name="Reporter",
            role="reporter",
            model=OpenAIChat(id="gpt-4o"),
            instructions=AGENT_INSTRUCTIONS,
        )

    def run(self) -> dict[str, Any]:
        self.logger.info("[Reporter] Starting report generation")

        try:
            # Collect all artifacts from team memory
            artifacts = self._collect_artifacts()

            # Generate executive summary using LLM
            executive_summary = self._generate_executive_summary(artifacts)

            # Generate markdown report
            md_content = self._generate_markdown_report(artifacts, executive_summary)
            md_path = self.run_dir / "report.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)
            self.logger.info(f"[Reporter] Markdown report saved to {md_path}")

            # Generate HTML report
            html_path = self._generate_html_report(md_content)

            # Generate DOCX report
            docx_path = self._generate_docx_report(md_content)

            # Generate Dashboard
            dashboard_path = self._generate_dashboard(artifacts, executive_summary)

            # Update team memory
            self.team_memory["report.md"] = str(md_path)
            self.team_memory["report.html"] = str(html_path) if html_path else None
            self.team_memory["report.docx"] = str(docx_path) if docx_path else None
            self.team_memory["dashboard.html"] = str(dashboard_path) if dashboard_path else None
            self.team_memory["executive_summary"] = executive_summary

            self.logger.info("[Reporter] Report generation complete")

            return {
                "status": "success",
                "markdown_path": str(md_path),
                "html_path": str(html_path) if html_path else None,
                "docx_path": str(docx_path) if docx_path else None,
                "dashboard_path": str(dashboard_path) if dashboard_path else None,
                "executive_summary": executive_summary,
            }

        except Exception as e:
            self.logger.error(f"[Reporter] Error during report generation: {e}", exc_info=True)
            return {
                "status": "error",
                "error": str(e),
            }

    def _collect_artifacts(self) -> dict[str, Any]:
        artifacts = {
            "timestamp": datetime.now().isoformat(),
            "run_dir": str(self.run_dir),
        }

        # EDA (now includes ingestion)
        if "eda_summary" in self.team_memory:
            artifacts["eda"] = self.team_memory["eda_summary"]

        # LLM insights from EDA
        if "eda_llm_analysis" in self.team_memory:
            artifacts["eda_insights"] = self.team_memory["eda_llm_analysis"]

        # Visualizations
        if "plots_index" in self.team_memory:
            artifacts["visualizations"] = self.team_memory["plots_index"]

        # LLM insights from visualizations
        if "viz_llm_insights" in self.team_memory:
            artifacts["viz_insights"] = self.team_memory["viz_llm_insights"]

        # Feature selection
        if "feature_ranking" in self.team_memory:
            artifacts["feature_selection"] = {
                "top_features": self.team_memory.get("feature_list", []),
                "ranking": self.team_memory["feature_ranking"][:20],  # Top 20
            }

        # LLM insights from feature selection
        if "feature_selection_insights" in self.team_memory:
            artifacts["feature_insights"] = self.team_memory["feature_selection_insights"]

        # Modeling
        if "models_summary" in self.team_memory:
            artifacts["modeling"] = self.team_memory["models_summary"]

        # LLM insights from modeling
        if "modeling_insights" in self.team_memory:
            artifacts["modeling_insights"] = self.team_memory["modeling_insights"]

        return artifacts

    def _generate_executive_summary(self, artifacts: dict[str, Any]) -> str:
        try:
            # Prepare consolidated information
            summary_parts = []

            # Data overview
            if "eda" in artifacts:
                eda = artifacts["eda"]
                summary_parts.append(
                    f"DATA: {eda.get('rows', 'N/A')} rows, {eda.get('columns', 'N/A')} columns"
                )

            # Feature selection
            if "feature_selection" in artifacts:
                fs = artifacts["feature_selection"]
                summary_parts.append(
                    f"FEATURES: Selected {len(fs['top_features'])} features from {len(fs.get('ranking', []))} total"
                )

            # Best model
            if "modeling" in artifacts:
                mod = artifacts["modeling"]
                summary_parts.append(
                    f"MODEL: {mod.get('best_model', 'N/A')} for {mod.get('task_type', 'N/A')} on '{mod.get('target_column', 'N/A')}'"
                )

            # Collect all LLM insights
            insights_collection = []
            if "eda_insights" in artifacts:
                insights_collection.append(f"EDA INSIGHTS:\n{artifacts['eda_insights']}")
            if "viz_insights" in artifacts:
                insights_collection.append(f"VISUALIZATION INSIGHTS:\n{artifacts['viz_insights']}")
            if "feature_insights" in artifacts:
                insights_collection.append(
                    f"FEATURE SELECTION INSIGHTS:\n{artifacts['feature_insights']}"
                )
            if "modeling_insights" in artifacts:
                insights_collection.append(f"MODELING INSIGHTS:\n{artifacts['modeling_insights']}")

            prompt = EXECUTIVE_SUMMARY_PROMPT_TEMPLATE.format(
                pipeline_overview=chr(10).join(summary_parts),
                stage_insights=chr(10).join(insights_collection),
            )

            self.logger.info("[Reporter] Generating executive summary with LLM")
            response = self.llm_agent.run(input=prompt)
            summary_text = response.content

            self.logger.info("[Reporter] Executive summary generated")
            return summary_text

        except Exception as e:
            self.logger.error(f"[Reporter] Error generating executive summary: {e}")
            return f"Error generating executive summary: {str(e)}"

    def _generate_markdown_report(self, artifacts: dict[str, Any], executive_summary: str) -> str:
        md = []

        # Header
        md.append("# RenaML POC - Automated ML Pipeline Report")
        md.append(f"\n**Generated:** {artifacts['timestamp']}")
        md.append(f"\n**Run Directory:** `{artifacts['run_dir']}`")
        md.append("\n---\n")

        # Executive Summary (NEW)
        md.append("## Executive Summary")
        md.append(f"\n{executive_summary}")
        md.append("\n---\n")

        # Table of Contents
        md.append("## Table of Contents")
        md.append("1. [Executive Summary](#executive-summary)")
        md.append("2. [Exploratory Data Analysis](#exploratory-data-analysis)")
        md.append("3. [Visualizations](#visualizations)")
        md.append("4. [Feature Selection](#feature-selection)")
        md.append("5. [Model Training](#model-training)")
        md.append("\n---\n")

        # EDA
        if "eda" in artifacts:
            md.append("## Exploratory Data Analysis")
            eda = artifacts["eda"]

            if "numeric_columns" in eda:
                md.append(f"\n### Numeric Features ({len(eda['numeric_columns'])})")
                md.append(f"\n{', '.join(eda['numeric_columns'][:10])}")
                if len(eda["numeric_columns"]) > 10:
                    md.append(f"... and {len(eda['numeric_columns']) - 10} more")

            if "categorical_columns" in eda:
                md.append(f"\n### Categorical Features ({len(eda['categorical_columns'])})")
                md.append(f"\n{', '.join(eda['categorical_columns'][:10])}")
                if len(eda["categorical_columns"]) > 10:
                    md.append(f"... and {len(eda['categorical_columns']) - 10} more")

            if "alerts" in eda and eda["alerts"]:
                md.append("\n### Alerts")
                for alert in eda["alerts"]:
                    md.append(f"- ⚠️ {alert}")

            md.append("\n")

        # Visualizations
        if "visualizations" in artifacts:
            md.append("## Visualizations")
            md.append(f"\n**Total plots generated:** {len(artifacts['visualizations'])}")
            for viz in artifacts["visualizations"]:
                md.append(f"\n- **{viz['type']}**: {viz['description']}")
                md.append(f"  - Path: `{viz['path']}`")
            md.append("\n")

        # Feature Selection
        if "feature_selection" in artifacts:
            md.append("## Feature Selection")
            fs = artifacts["feature_selection"]
            md.append(f"\n**Top {len(fs['top_features'])} Selected Features:**")
            md.append("\n| Rank | Feature | Score |")
            md.append("|------|---------|-------|")
            for i, (feat, score) in enumerate(fs["ranking"][:20], 1):
                md.append(f"| {i} | {feat} | {score:.4f} |")
            md.append("\n")

        # Model Training
        if "modeling" in artifacts:
            md.append("## Model Training")
            mod = artifacts["modeling"]
            md.append(f"\n- **Task Type:** {mod.get('task_type', 'N/A')}")
            md.append(f"- **Target Column:** {mod.get('target_column', 'N/A')}")
            md.append(f"- **Best Model:** {mod.get('best_model', 'N/A')}")
            md.append(f"- **Model Path:** `{mod.get('model_path', 'N/A')}`")

            if "comparison_results" in mod and mod["comparison_results"]:
                md.append("\n### Model Comparison Results")
                md.append("\nTop 5 models:")
                for i, result in enumerate(mod["comparison_results"][:5], 1):
                    md.append(f"\n{i}. **{result.get('Model', 'Unknown')}**")
                    # Show first few metrics
                    for key, value in list(result.items())[:5]:
                        if key != "Model":
                            md.append(f"   - {key}: {value}")

            md.append("\n")

        # Footer
        md.append("---")
        md.append("\n*Report generated by RenaML POC - Automated ML Pipeline*")

        return "\n".join(md)

    def _generate_html_report(self, md_content: str) -> Path | None:
        try:
            html_content = markdown.markdown(
                md_content, extensions=["tables", "fenced_code", "toc"]
            )

            # Add basic styling
            html_template = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>RenaML POC Report</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #4CAF50;
            color: white;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
        }}
        h1, h2, h3 {{
            color: #333;
        }}
        hr {{
            border: none;
            border-top: 2px solid #ddd;
            margin: 30px 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""

            html_path = self.run_dir / "report.html"
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_template)

            self.logger.info(f"[Reporter] HTML report saved to {html_path}")
            return html_path

        except Exception as e:
            self.logger.error(f"[Reporter] Error generating HTML report: {e}")
            return None

    def _generate_docx_report(self, md_content: str) -> Path | None:
        try:
            doc = Document()

            # Add title
            title = doc.add_heading("RenaML POC - Automated ML Pipeline Report", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Parse markdown and add content
            lines = md_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("|"):
                    # Skip table formatting for now (complex)
                    continue
                elif line.startswith("---"):
                    doc.add_paragraph("_" * 50)
                else:
                    doc.add_paragraph(line)

            docx_path = self.run_dir / "report.docx"
            doc.save(docx_path)

            self.logger.info(f"[Reporter] DOCX report saved to {docx_path}")
            return docx_path

        except ImportError:
            self.logger.warning("[Reporter] python-docx not available, skipping DOCX generation")
            return None
        except Exception as e:
            self.logger.error(f"[Reporter] Error generating DOCX report: {e}")
            return None

    def _generate_dashboard(self, artifacts: dict[str, Any], executive_summary: str) -> Path | None:
        """Generate interactive Bootstrap dashboard with embedded visualizations."""
        try:
            # Extract data from artifacts
            eda = artifacts.get("eda", {})
            visualizations = artifacts.get("visualizations", [])
            feature_selection = artifacts.get("feature_selection", {})
            modeling = artifacts.get("modeling", {})
            
            # Load feature selection details if available
            feature_methods = {}
            feature_ranking = feature_selection.get("ranking", [])
            
            # Try to load detailed feature selection data
            feature_json_path = self.run_dir / "feature_selection" / "feature_selection_report.json"
            if feature_json_path.exists():
                import json
                with open(feature_json_path, "r") as f:
                    fs_data = json.load(f)
                    feature_methods = fs_data.get("methods", {})
            
            # Prepare feature selection table rows
            feature_rows = []
            for i, (feat, score) in enumerate(feature_ranking[:10], 1):
                badge_colors = ["warning", "secondary", "info", "light text-dark", "light text-dark",
                               "light text-dark", "light text-dark", "light text-dark", "light text-dark", "light text-dark"]
                badge_color = badge_colors[min(i-1, len(badge_colors)-1)]
                
                # Get method-specific scores
                stat_score = feature_methods.get("statistical", {}).get(feat, 0)
                tree_score = feature_methods.get("tree_based", {}).get(feat, 0)
                mi_score = feature_methods.get("mutual_info", {}).get(feat, 0)
                
                feature_rows.append(f"""
                                            <tr>
                                                <td><span class="badge bg-{badge_color} badge-rank">{i}º</span></td>
                                                <td><strong>{feat}</strong></td>
                                                <td>{score:.2f}</td>
                                                <td>{stat_score:.2f}</td>
                                                <td>{tree_score:.3f}</td>
                                                <td>{mi_score:.3f}</td>
                                            </tr>""")
            
            feature_table_html = "\n".join(feature_rows) if feature_rows else """
                                            <tr>
                                                <td colspan="6" class="text-center text-muted">Nenhuma feature selecionada</td>
                                            </tr>"""
            
            # Count features
            num_features = len(eda.get("numeric_columns", []))
            cat_features = len(eda.get("categorical_columns", []))
            num_viz = len(visualizations)
            task_type = modeling.get("task_type", "N/A")
            
            # Generate timestamp
            timestamp = datetime.now().strftime("%d de %B, %Y às %H:%M:%S")
            timestamp_short = datetime.now().strftime("%H:%M:%S")
            
            # Prepare modeling status
            if modeling.get("error"):
                modeling_alert = f"""
                                <div class="alert alert-warning">
                                    <i class="bi bi-exclamation-triangle-fill"></i> 
                                    <strong>Atenção:</strong> Ocorreu um erro durante o processo de modelagem.
                                    <br>
                                    <code>{modeling.get('error', 'Unknown error')}</code>
                                </div>"""
                modeling_metrics = """
                                    <div class="col-md-4">
                                        <div class="metric-box">
                                            <div class="metric-title">Task Type</div>
                                            <div class="metric-value">N/A</div>
                                        </div>
                                    </div>
                                    <div class="col-md-4">
                                        <div class="metric-box">
                                            <div class="metric-title">Target Column</div>
                                            <div class="metric-value">N/A</div>
                                        </div>
                                    </div>
                                    <div class="col-md-4">
                                        <div class="metric-box">
                                            <div class="metric-title">Best Model</div>
                                            <div class="metric-value">N/A</div>
                                        </div>
                                    </div>"""
            else:
                modeling_alert = f"""
                                <div class="alert alert-success">
                                    <i class="bi bi-check-circle-fill"></i> 
                                    <strong>Sucesso!</strong> Modelagem concluída com sucesso.
                                </div>"""
                modeling_metrics = f"""
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Task Type</div>
                                            <div class="metric-value">{modeling.get('task_type', 'N/A')}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Target Column</div>
                                            <div class="metric-value">{modeling.get('target_column', 'N/A')}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Best Model</div>
                                            <div class="metric-value">{modeling.get('best_model', 'N/A')}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Model Score</div>
                                            <div class="metric-value">{modeling.get('best_score', 'N/A')}</div>
                                        </div>
                                    </div>"""
            
            # List numeric and categorical features
            numeric_features_list = "\n".join([f'<li class="list-group-item">{f}</li>' 
                                               for f in eda.get("numeric_columns", [])[:10]])
            if not numeric_features_list:
                numeric_features_list = '<li class="list-group-item text-muted">Nenhuma feature numérica</li>'
            
            categorical_features_list = "\n".join([f'<li class="list-group-item">{f}</li>' 
                                                   for f in eda.get("categorical_columns", [])[:10]])
            if not categorical_features_list:
                categorical_features_list = '<li class="list-group-item text-muted">Nenhuma feature categórica</li>'
            
            # Target variable badge
            target_col = modeling.get("target_column", "N/A")
            if target_col != "N/A":
                categorical_features_list = categorical_features_list.replace(
                    f'<li class="list-group-item">{target_col}</li>',
                    f'<li class="list-group-item"><strong>{target_col}</strong><span class="badge bg-info float-end">Target Variable</span></li>'
                )

            dashboard_html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>RenaML - Dashboard de Análise de Machine Learning</title>
    
    <!-- Bootstrap CSS -->
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
    <!-- Bootstrap Icons -->
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">
    <!-- Plotly -->
    <script src="https://cdn.plot.ly/plotly-3.2.0.min.js"></script>
    
    <style>
        :root {{
            --primary-color: #2c3e50;
            --secondary-color: #3498db;
            --success-color: #27ae60;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --info-color: #16a085;
            --light-bg: #ecf0f1;
            --dark-text: #2c3e50;
        }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f8f9fa;
        }}
        
        .navbar {{
            background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
            box-shadow: 0 2px 4px rgba(0,0,0,.1);
        }}
        
        .navbar-brand {{
            font-weight: 700;
            font-size: 1.5rem;
            color: white !important;
        }}
        
        .main-container {{
            margin-top: 20px;
            margin-bottom: 40px;
        }}
        
        .card {{
            border: none;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0,0,0,.07);
            margin-bottom: 20px;
            transition: transform 0.2s;
        }}
        
        .card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 8px 12px rgba(0,0,0,.15);
        }}
        
        .card-header {{
            background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
            color: white;
            font-weight: 600;
            border-radius: 10px 10px 0 0;
            padding: 1rem 1.5rem;
        }}
        
        .nav-tabs {{
            border-bottom: 2px solid var(--secondary-color);
        }}
        
        .nav-tabs .nav-link {{
            color: var(--dark-text);
            border: none;
            padding: 1rem 1.5rem;
            font-weight: 500;
            transition: all 0.3s;
        }}
        
        .nav-tabs .nav-link:hover {{
            border: none;
            color: var(--secondary-color);
            background-color: rgba(52, 152, 219, 0.1);
        }}
        
        .nav-tabs .nav-link.active {{
            color: var(--secondary-color);
            background-color: white;
            border: none;
            border-bottom: 3px solid var(--secondary-color);
        }}
        
        .stat-card {{
            text-align: center;
            padding: 1.5rem;
            background: white;
            border-radius: 10px;
            height: 100%;
        }}
        
        .stat-icon {{
            font-size: 2.5rem;
            margin-bottom: 0.5rem;
        }}
        
        .stat-value {{
            font-size: 2rem;
            font-weight: 700;
            margin: 0.5rem 0;
        }}
        
        .stat-label {{
            color: #6c757d;
            font-size: 0.9rem;
        }}
        
        .table-responsive {{
            border-radius: 10px;
            overflow: hidden;
        }}
        
        .table {{
            margin-bottom: 0;
        }}
        
        .table thead {{
            background: linear-gradient(135deg, var(--primary-color) 0%, var(--secondary-color) 100%);
            color: white;
        }}
        
        .table-striped tbody tr:nth-of-type(odd) {{
            background-color: rgba(52, 152, 219, 0.05);
        }}
        
        .badge-rank {{
            font-size: 1rem;
            padding: 0.5rem 0.75rem;
        }}
        
        .plot-container {{
            background: white;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        
        .alert-custom {{
            border-left: 4px solid;
            border-radius: 5px;
        }}
        
        .alert-info-custom {{
            border-left-color: var(--info-color);
            background-color: rgba(22, 160, 133, 0.1);
        }}
        
        .footer {{
            background-color: var(--primary-color);
            color: white;
            padding: 2rem 0;
            margin-top: 3rem;
        }}
        
        .tab-content {{
            padding: 2rem 0;
        }}
        
        iframe {{
            width: 100%;
            border: none;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0,0,0,.1);
        }}
        
        .metric-box {{
            padding: 1rem;
            background: linear-gradient(135deg, rgba(52, 152, 219, 0.1) 0%, rgba(41, 128, 185, 0.1) 100%);
            border-radius: 8px;
            margin-bottom: 1rem;
        }}
        
        .metric-title {{
            font-weight: 600;
            color: var(--secondary-color);
            margin-bottom: 0.5rem;
        }}
        
        .metric-value {{
            font-size: 1.25rem;
            font-weight: 700;
            color: var(--primary-color);
        }}
    </style>
</head>
<body>
    <!-- Navbar -->
    <nav class="navbar navbar-expand-lg navbar-dark">
        <div class="container-fluid">
            <a class="navbar-brand" href="#">
                <i class="bi bi-graph-up-arrow"></i> RenaML Dashboard
            </a>
            <span class="navbar-text text-white">
                <i class="bi bi-calendar3"></i> {timestamp}
            </span>
        </div>
    </nav>

    <!-- Main Container -->
    <div class="container-fluid main-container">
        <!-- Header Stats -->
        <div class="row mb-4">
            <div class="col-md-3 mb-3">
                <div class="card stat-card">
                    <div class="stat-icon text-primary">
                        <i class="bi bi-bar-chart-fill"></i>
                    </div>
                    <div class="stat-value text-primary">{num_features}</div>
                    <div class="stat-label">Features Numéricas</div>
                </div>
            </div>
            <div class="col-md-3 mb-3">
                <div class="card stat-card">
                    <div class="stat-icon text-success">
                        <i class="bi bi-tags-fill"></i>
                    </div>
                    <div class="stat-value text-success">{cat_features}</div>
                    <div class="stat-label">Features Categóricas</div>
                </div>
            </div>
            <div class="col-md-3 mb-3">
                <div class="card stat-card">
                    <div class="stat-icon text-warning">
                        <i class="bi bi-diagram-3-fill"></i>
                    </div>
                    <div class="stat-value text-warning">{num_viz}</div>
                    <div class="stat-label">Visualizações</div>
                </div>
            </div>
            <div class="col-md-3 mb-3">
                <div class="card stat-card">
                    <div class="stat-icon text-info">
                        <i class="bi bi-cpu-fill"></i>
                    </div>
                    <div class="stat-value text-info" style="font-size: 1.5rem;">{task_type}</div>
                    <div class="stat-label">Tipo de Tarefa</div>
                </div>
            </div>
        </div>

        <!-- Tabs Navigation -->
        <ul class="nav nav-tabs" id="analysisTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="overview-tab" data-bs-toggle="tab" data-bs-target="#overview" type="button" role="tab">
                    <i class="bi bi-house-fill"></i> Visão Geral
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="eda-tab" data-bs-toggle="tab" data-bs-target="#eda" type="button" role="tab">
                    <i class="bi bi-clipboard-data"></i> Análise Exploratória
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="viz-tab" data-bs-toggle="tab" data-bs-target="#viz" type="button" role="tab">
                    <i class="bi bi-graph-up"></i> Visualizações
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="features-tab" data-bs-toggle="tab" data-bs-target="#features" type="button" role="tab">
                    <i class="bi bi-list-check"></i> Seleção de Features
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="models-tab" data-bs-toggle="tab" data-bs-target="#models" type="button" role="tab">
                    <i class="bi bi-robot"></i> Modelagem
                </button>
            </li>
        </ul>

        <!-- Tabs Content -->
        <div class="tab-content" id="analysisTabsContent">
            <!-- Overview Tab -->
            <div class="tab-pane fade show active" id="overview" role="tabpanel">
                <div class="row">
                    <div class="col-md-12">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-info-circle-fill"></i> Resumo Executivo
                            </div>
                            <div class="card-body">
                                <h5>RenaML POC - Automated ML Pipeline Report</h5>
                                <p class="text-muted">
                                    <i class="bi bi-clock"></i> Gerado em: {timestamp}
                                </p>
                                <p class="text-muted">
                                    <i class="bi bi-folder"></i> Diretório: <code>{artifacts.get('run_dir', 'N/A')}</code>
                                </p>
                                
                                <div class="alert alert-custom alert-info-custom mt-4">
                                    <h6><i class="bi bi-lightbulb-fill"></i> Executive Summary</h6>
                                    <div style="white-space: pre-wrap;">{executive_summary}</div>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>

                <div class="row mt-4">
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-list-ul"></i> Features Numéricas
                            </div>
                            <div class="card-body">
                                <div class="metric-box">
                                    <div class="metric-title">Total de Features</div>
                                    <div class="metric-value">{num_features}</div>
                                </div>
                                <ul class="list-group list-group-flush">
{numeric_features_list}
                                </ul>
                            </div>
                        </div>
                    </div>

                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-tags"></i> Features Categóricas
                            </div>
                            <div class="card-body">
                                <div class="metric-box">
                                    <div class="metric-title">Total de Features</div>
                                    <div class="metric-value">{cat_features}</div>
                                </div>
                                <ul class="list-group list-group-flush">
{categorical_features_list}
                                </ul>
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <!-- EDA Tab -->
            <div class="tab-pane fade" id="eda" role="tabpanel">
                <div class="row">
                    <div class="col-md-12">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-file-earmark-bar-graph"></i> Relatórios de Análise Exploratória
                            </div>
                            <div class="card-body">
                                <div class="alert alert-custom alert-info-custom">
                                    <i class="bi bi-info-circle"></i> 
                                    Os relatórios abaixo foram gerados usando bibliotecas especializadas de profiling de dados.
                                </div>
                                
                                <div class="row">
                                    <div class="col-md-6 mb-3">
                                        <div class="card h-100">
                                            <div class="card-body text-center">
                                                <i class="bi bi-file-earmark-text display-4 text-primary mb-3"></i>
                                                <h5 class="card-title">Sweetviz Report</h5>
                                                <p class="card-text">Análise visual comparativa dos dados</p>
                                                <a href="eda/sweetviz_report.html" target="_blank" class="btn btn-primary">
                                                    <i class="bi bi-box-arrow-up-right"></i> Abrir Relatório
                                                </a>
                                            </div>
                                        </div>
                                    </div>
                                    <div class="col-md-6 mb-3">
                                        <div class="card h-100">
                                            <div class="card-body text-center">
                                                <i class="bi bi-file-earmark-bar-graph display-4 text-success mb-3"></i>
                                                <h5 class="card-title">YData Profile Report</h5>
                                                <p class="card-text">Análise detalhada de qualidade de dados</p>
                                                <a href="eda/ydata_profile.html" target="_blank" class="btn btn-success">
                                                    <i class="bi bi-box-arrow-up-right"></i> Abrir Relatório
                                                </a>
                                            </div>
                                        </div>
                                    </div>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Visualizations Tab -->
            <div class="tab-pane fade" id="viz" role="tabpanel">
                <div class="row">
                    <div class="col-md-12 mb-4">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-bar-chart-line"></i> Distribuições de Features Numéricas
                            </div>
                            <div class="card-body">
                                <iframe src="visualizations/distributions.html" height="650"></iframe>
                            </div>
                        </div>
                    </div>

                    <div class="col-md-12 mb-4">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-grid-3x3"></i> Matriz de Correlação
                            </div>
                            <div class="card-body">
                                <iframe src="visualizations/correlation_matrix.html" height="850"></iframe>
                            </div>
                        </div>
                    </div>

                    <div class="col-md-12 mb-4">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-pie-chart"></i> Distribuições de Features Categóricas
                            </div>
                            <div class="card-body">
                                <iframe src="visualizations/categorical_distributions.html" height="450"></iframe>
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Feature Selection Tab -->
            <div class="tab-pane fade" id="features" role="tabpanel">
                <div class="row">
                    <div class="col-md-12">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-filter-circle"></i> Ranking de Features
                            </div>
                            <div class="card-body">
                                <div class="row mb-4">
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Total de Features</div>
                                            <div class="metric-value">{len(feature_ranking)}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Target</div>
                                            <div class="metric-value">{target_col}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Tipo de Tarefa</div>
                                            <div class="metric-value">{task_type}</div>
                                        </div>
                                    </div>
                                    <div class="col-md-3">
                                        <div class="metric-box">
                                            <div class="metric-title">Timestamp</div>
                                            <div class="metric-value" style="font-size: 0.9rem;">{timestamp_short}</div>
                                        </div>
                                    </div>
                                </div>

                                <h5 class="mb-3"><i class="bi bi-trophy-fill text-warning"></i> Top Features Selecionadas</h5>
                                <div class="table-responsive">
                                    <table class="table table-striped table-hover">
                                        <thead>
                                            <tr>
                                                <th>Rank</th>
                                                <th>Feature</th>
                                                <th>Score Combinado</th>
                                                <th>Score Estatístico</th>
                                                <th>Score Tree-Based</th>
                                                <th>Mutual Information</th>
                                            </tr>
                                        </thead>
                                        <tbody>{feature_table_html}
                                        </tbody>
                                    </table>
                                </div>

                                <div class="alert alert-custom alert-info-custom mt-4">
                                    <h6><i class="bi bi-lightbulb"></i> Interpretação</h6>
                                    <p class="mb-0">
                                        O ranking de features foi calculado usando três métodos complementares: 
                                        análise estatística (ANOVA/Chi-squared), importância baseada em árvores (Random Forest), 
                                        e informação mútua. As features com menores scores combinados são consideradas mais relevantes.
                                    </p>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Models Tab -->
            <div class="tab-pane fade" id="models" role="tabpanel">
                <div class="row">
                    <div class="col-md-12">
                        <div class="card">
                            <div class="card-header">
                                <i class="bi bi-cpu"></i> Resultados de Modelagem
                            </div>
                            <div class="card-body">
{modeling_alert}

                                <div class="row">
{modeling_metrics}
                                </div>

                                <div class="alert alert-custom alert-info-custom mt-4">
                                    <h6><i class="bi bi-info-circle"></i> Sobre a Modelagem</h6>
                                    <p class="mb-0">
                                        O pipeline RenaML utiliza AutoML para treinar e comparar múltiplos modelos automaticamente,
                                        selecionando o melhor com base em métricas de performance. Para análise detalhada dos resultados,
                                        consulte os arquivos JSON na pasta de modelagem.
                                    </p>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- Footer -->
    <footer class="footer">
        <div class="container text-center">
            <p class="mb-0">
                <i class="bi bi-robot"></i> RenaML POC - Automated ML Pipeline
            </p>
            <p class="text-white-50 mb-0">
                Dashboard gerado automaticamente em {timestamp}
            </p>
        </div>
    </footer>

    <!-- Bootstrap JS -->
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/js/bootstrap.bundle.min.js"></script>
    
    <script>
        // Add smooth scrolling
        document.querySelectorAll('a[href^="#"]').forEach(anchor => {{
            anchor.addEventListener('click', function (e) {{
                e.preventDefault();
                const target = this.getAttribute('href');
                if (target !== '#') {{
                    document.querySelector(target).scrollIntoView({{
                        behavior: 'smooth'
                    }});
                }}
            }});
        }});

        // Tab change animation
        document.querySelectorAll('[data-bs-toggle="tab"]').forEach(tab => {{
            tab.addEventListener('shown.bs.tab', function (e) {{
                const targetPane = document.querySelector(e.target.getAttribute('data-bs-target'));
                targetPane.style.opacity = '0';
                setTimeout(() => {{
                    targetPane.style.transition = 'opacity 0.3s';
                    targetPane.style.opacity = '1';
                }}, 10);
            }});
        }});

        // Console welcome message
        console.log('%c RenaML Dashboard ', 'background: #2c3e50; color: #fff; font-size: 20px; padding: 10px;');
        console.log('%c Powered by Bootstrap 5 & Plotly ', 'background: #3498db; color: #fff; font-size: 14px; padding: 5px;');
    </script>
</body>
</html>"""

            dashboard_path = self.run_dir / "dashboard.html"
            with open(dashboard_path, "w", encoding="utf-8") as f:
                f.write(dashboard_html)

            self.logger.info(f"[Reporter] Dashboard saved to {dashboard_path}")
            return dashboard_path

        except Exception as e:
            self.logger.error(f"[Reporter] Error generating dashboard: {e}", exc_info=True)
            return None
